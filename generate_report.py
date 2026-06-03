"""
Genera el reporte del Proyecto 3 en formato .docx.
Ejecutar una sola vez: python generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from tarea3 import build_sequence as min_seq
from tarea4 import build_sequence as worst_seq
from algorithms import run_mtf, run_imtf

# ── Colores ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x37, 0x64)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF2, 0xF2, 0xF2)

CONFIG = [0, 1, 2, 3, 4]

# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _bold_run(para, text: str, size: int = 10):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return run


def _heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = NAVY
    return p


def _mono(doc, text: str):
    """Parrafo con fuente monoespaciada para mostrar tablas de consola."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p


def _spacer(doc):
    doc.add_paragraph()


def _mtf_table(doc, config, requests):
    """Renderiza la ejecucion MTF como tabla Word."""
    lst   = list(config)
    rows  = []
    total = 0
    for req in requests:
        before = list(lst)
        pos    = lst.index(req) + 1
        total += pos
        lst.remove(req)
        lst.insert(0, req)
        rows.append((str(before), req, pos, str(lst)))

    tbl = doc.add_table(rows=1 + len(rows) + 1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Encabezado
    hdr_cells = tbl.rows[0].cells
    for i, txt in enumerate(["Config antes", "Solicitud", "Costo", "Config despues"]):
        hdr_cells[i].text = txt
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        _set_cell_bg(hdr_cells[i], "1F3764")

    # Datos
    for r_idx, (before, req, cost, after) in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        cells[0].text = before
        cells[1].text = str(req)
        cells[2].text = str(cost)
        cells[3].text = after
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c in cells:
            c.paragraphs[0].runs[0].font.size = Pt(8)
            c.paragraphs[0].runs[0].font.name = "Courier New"
            _set_cell_bg(c, bg)

    # Total
    total_row = tbl.rows[-1].cells
    total_row[0].merge(total_row[1]).merge(total_row[2])
    total_row[0].text = "Costo total de acceso"
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[-1].text = str(total)
    total_row[-1].paragraphs[0].runs[0].bold = True
    _set_cell_bg(total_row[0], "D9E1F2")
    _set_cell_bg(total_row[-1], "D9E1F2")

    return total


def _imtf_table(doc, config, requests):
    """Renderiza la ejecucion IMTF como tabla Word."""
    lst   = list(config)
    rows  = []
    total = 0
    for idx, req in enumerate(requests):
        before = list(lst)
        pos    = lst.index(req) + 1
        total += pos
        window = requests[idx + 1 : idx + pos]
        move   = req in window
        if move:
            lst.remove(req)
            lst.insert(0, req)
        rows.append((str(before), req, pos, "Si" if move else "No", str(lst)))

    tbl = doc.add_table(rows=1 + len(rows) + 1, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = tbl.rows[0].cells
    for i, txt in enumerate(["Config antes", "Solicitud", "Costo", "Mueve?", "Config despues"]):
        hdr_cells[i].text = txt
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        _set_cell_bg(hdr_cells[i], "1F3764")

    for r_idx, (before, req, cost, move, after) in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        cells[0].text = before
        cells[1].text = str(req)
        cells[2].text = str(cost)
        cells[3].text = move
        cells[4].text = after
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c in cells:
            c.paragraphs[0].runs[0].font.size = Pt(8)
            c.paragraphs[0].runs[0].font.name = "Courier New"
            _set_cell_bg(c, bg)

    total_row = tbl.rows[-1].cells
    total_row[0].merge(total_row[1]).merge(total_row[2]).merge(total_row[3])
    total_row[0].text = "Costo total de acceso"
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[-1].text = str(total)
    total_row[-1].paragraphs[0].runs[0].bold = True
    _set_cell_bg(total_row[0], "D9E1F2")
    _set_cell_bg(total_row[-1], "D9E1F2")

    return total


# ── Documento ─────────────────────────────────────────────────────────────────

def build_report():
    doc = Document()

    # Margenes
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Portada ──────────────────────────────────────────────────────────────
    title = doc.add_heading("Proyecto No. 3", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = NAVY

    sub = doc.add_paragraph("Analisis y Diseno de Algoritmos")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)

    sub2 = doc.add_paragraph("Seccion 10 — Gabriel Brolo | 2026")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.runs[0].font.size = Pt(11)
    sub2.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    doc.add_page_break()

    # ── Enlace a repositorio y video ─────────────────────────────────────────
    _heading(doc, "1. Repositorio y Video", level=1)
    p = doc.add_paragraph()
    p.add_run("Repositorio GitHub: ").bold = True
    p.add_run("[enlace pendiente de configurar]")

    p2 = doc.add_paragraph()
    p2.add_run("Video YouTube:      ").bold = True
    p2.add_run("[enlace pendiente de subir]")
    _spacer(doc)

    # ── Tarea 1 ──────────────────────────────────────────────────────────────
    _heading(doc, "2. Tarea 1 — MTF: secuencia 0 1 2 3 4 repetida", level=1)
    doc.add_paragraph(
        "Configuracion inicial: [0, 1, 2, 3, 4]\n"
        "Secuencia de solicitudes: 0 1 2 3 4 0 1 2 3 4 0 1 2 3 4 0 1 2 3 4 (20 elementos)"
    )
    t1_cost = _mtf_table(doc, CONFIG,
                         [0,1,2,3,4, 0,1,2,3,4, 0,1,2,3,4, 0,1,2,3,4])
    _spacer(doc)
    p = doc.add_paragraph()
    _bold_run(p, "Analisis: ")
    p.add_run(
        "El primer recorrido (0→4) cuesta 1+2+3+4+5 = 15. "
        "Tras el primer pase la lista queda invertida [4,3,2,1,0]; "
        "los tres pases siguientes siempre encuentran cada elemento en la ultima posicion "
        "(costo 5 por acceso), sumando 5x5x3 = 75. Total: 15 + 75 = 90."
    )
    _spacer(doc)

    # ── Tarea 2 ──────────────────────────────────────────────────────────────
    _heading(doc, "3. Tarea 2 — MTF: secuencia mixta", level=1)
    doc.add_paragraph(
        "Configuracion inicial: [0, 1, 2, 3, 4]\n"
        "Secuencia de solicitudes: 4 3 2 1 0 1 2 3 4 3 2 1 0 1 2 3 4 (17 elementos)"
    )
    _mtf_table(doc, CONFIG, [4,3,2,1,0, 1,2,3,4,3, 2,1,0,1,2, 3,4])
    _spacer(doc)

    # ── Tarea 3 ──────────────────────────────────────────────────────────────
    _heading(doc, "4. Tarea 3 — Secuencia de minimo costo (20 solicitudes)", level=1)
    min_seq_list = min_seq(CONFIG, 20)
    doc.add_paragraph(
        f"Secuencia de minimo costo: {min_seq_list}\n"
        f"Costo total minimo: 20"
    )
    _mtf_table(doc, CONFIG, min_seq_list)
    _spacer(doc)
    p = doc.add_paragraph()
    _bold_run(p, "Justificacion: ")
    p.add_run(
        "El elemento '0' ocupa la posicion 1 desde el inicio, por lo que cada acceso "
        "tiene costo 1. Ningun acceso puede costar menos de 1 (cota inferior absoluta). "
        "Para 20 solicitudes: 20 x 1 = 20. La secuencia [0]*20 alcanza esa cota."
    )
    _spacer(doc)

    # ── Tarea 4 ──────────────────────────────────────────────────────────────
    _heading(doc, "5. Tarea 4 — Secuencia de peor caso (20 solicitudes)", level=1)
    worst_seq_list = worst_seq(CONFIG, 20)
    doc.add_paragraph(
        f"Secuencia de peor caso: {worst_seq_list}\n"
        f"Costo total maximo: 100"
    )
    _mtf_table(doc, CONFIG, worst_seq_list)
    _spacer(doc)
    p = doc.add_paragraph()
    _bold_run(p, "Justificacion: ")
    p.add_run(
        "Con 5 elementos el costo maximo por acceso es 5. "
        "El ciclo [4,3,2,1,0] garantiza que el elemento buscado siempre este "
        "en la ultima posicion antes de ser accedido, y MTF lo rota al frente "
        "preparando el siguiente ciclo. Para 20 solicitudes: 20 x 5 = 100. "
        "Esta es la cota superior absoluta y la secuencia la alcanza."
    )
    _spacer(doc)

    # ── Tarea 5 ──────────────────────────────────────────────────────────────
    _heading(doc, "6. Tarea 5 — Elemento repetido 20 veces", level=1)

    doc.add_paragraph("Secuencia: elemento 2 repetido 20 veces")
    _mtf_table(doc, CONFIG, [2] * 20)
    _spacer(doc)

    doc.add_paragraph("Secuencia: elemento 3 repetido 20 veces")
    _mtf_table(doc, CONFIG, [3] * 20)
    _spacer(doc)

    p = doc.add_paragraph()
    _bold_run(p, "Patron observado: ")
    p.add_run(
        "Sea e el elemento repetido en posicion inicial p (1-indexada) "
        "y n la cantidad de repeticiones:\n\n"
        "    Costo total = p + (n - 1) x 1 = p + n - 1\n\n"
        "El primer acceso cuesta p (posicion inicial del elemento). "
        "MTF lo mueve al frente; todos los accesos siguientes cuestan 1.\n\n"
        "Elemento 2 (p=3): 3 + 19 = 22   |   Elemento 3 (p=4): 4 + 19 = 23\n\n"
        "Para cualquier elemento con repeticion de 20 veces, el costo solo depende "
        "de su posicion inicial. A mayor posicion inicial, mayor costo total, "
        "pero la diferencia es siempre 1 unidad por posicion adicional."
    )
    _spacer(doc)

    # ── Tarea 6 ──────────────────────────────────────────────────────────────
    _heading(doc, "7. Tarea 6 — IMTF sobre mejor y peor caso de MTF", level=1)
    doc.add_paragraph(
        "IMTF (Improved MTF, Mohanty & Tripathy): al acceder al elemento en "
        "posicion i, se mueve al frente SOLO SI aparece en los proximos i-1 "
        "elementos de la secuencia (look-ahead)."
    )

    _heading(doc, "IMTF sobre la secuencia de minimo costo de MTF", level=2)
    doc.add_paragraph(f"Secuencia: {min_seq_list}")
    imtf_min = _imtf_table(doc, CONFIG, min_seq_list)
    _spacer(doc)

    _heading(doc, "IMTF sobre la secuencia de maximo costo de MTF", level=2)
    doc.add_paragraph(f"Secuencia: {worst_seq_list}")
    imtf_worst = _imtf_table(doc, CONFIG, worst_seq_list)
    _spacer(doc)

    # Tabla resumen comparativo
    _heading(doc, "Resumen comparativo MTF vs IMTF", level=2)
    tbl = doc.add_table(rows=3, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Escenario", "MTF", "IMTF"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        _set_cell_bg(hdr[i], "1F3764")
    data = [
        ("Mejor caso (minimo costo MTF)", "20",  str(imtf_min)),
        ("Peor caso  (maximo costo MTF)", "100", str(imtf_worst)),
    ]
    for r_idx, (esc, mtf, imtf) in enumerate(data):
        row = tbl.rows[r_idx + 1].cells
        row[0].text = esc
        row[1].text = mtf
        row[2].text = imtf
        bg = "F2F2F2" if r_idx % 2 == 0 else "FFFFFF"
        for c in row:
            _set_cell_bg(c, bg)

    _spacer(doc)
    p = doc.add_paragraph()
    _bold_run(p, "Analisis: ")
    p.add_run(
        "Mejor caso: IMTF == MTF (20). El elemento '0' esta en la posicion 1; "
        "la ventana look-ahead tiene tamano i-1 = 0, por lo que nunca se activa "
        "el movimiento. El costo es identico al de MTF.\n\n"
        "Peor caso: IMTF (60) < MTF (100). Con la secuencia [4,3,2,1,0]*4, "
        "el look-ahead de cada acceso no contiene el elemento accedido "
        "(ej.: acceder al 4 con ventana [3,2,1,0]). "
        "Como ningun elemento se mueve, la lista permanece estatica [0,1,2,3,4] "
        "durante toda la secuencia. El costo por ciclo es 5+4+3+2+1 = 15; "
        "con 4 ciclos: 4 x 15 = 60. "
        "IMTF reduce el costo en un 40% frente a MTF en este caso."
    )

    # ── Guardar ──────────────────────────────────────────────────────────────
    path = "Reporte_Proyecto3_MTF.docx"
    doc.save(path)
    print(f"Reporte guardado: {path}")


if __name__ == "__main__":
    build_report()
